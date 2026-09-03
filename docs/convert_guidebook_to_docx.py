from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "GUIDEBOOK_ADMIN_CUSTOMER_20260821.md"
OUTPUT = BASE_DIR / "GUIDEBOOK_ADMIN_CUSTOMER_20260821.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def set_run_font(run, name="Calibri", size=11, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=10, color="0B2545" if bold else "1F2937", bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=CELL_MARGIN_DXA, bottom=80, end=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def parse_table(lines, start):
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|") and lines[idx].strip().endswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def column_widths(rows):
    col_count = max(len(row) for row in rows)
    weights = [8] * col_count
    for row in rows:
        for i, cell in enumerate(row):
            weights[i] = max(weights[i], min(len(cell), 42))
    total = sum(weights)
    widths = [max(1200, int(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    delta = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += delta
    return widths


def add_markdown_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=col_count)
    table.style = "Table Grid"
    widths = column_widths(normalized)
    set_table_geometry(table, widths)

    for row_idx, row in enumerate(normalized):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            is_header = row_idx == 0
            set_cell_text(cell, value, bold=is_header)
            if is_header:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("KIU Store - Guidebook Pengguna")
        set_run_font(run, size=9, color="6B7280")


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Guidebook Pengguna KIU Store")
    set_run_font(run, size=22, color="0B2545", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Panduan bahasa umum untuk pelanggan, tim internal, dan pihak terkait.")
    set_run_font(run, size=11, color="555555")


def add_paragraph_from_text(doc, text):
    stripped = text.strip()
    if not stripped:
        return

    heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 3)
        doc.add_heading(heading.group(2).strip(), level=level)
        return

    if re.match(r"^\d+\.\s+", stripped):
        paragraph = doc.add_paragraph(style="List Number")
        content = re.sub(r"^\d+\.\s+", "", stripped)
        run = paragraph.add_run(content)
        set_run_font(run)
        return

    if stripped.startswith("- "):
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(stripped[2:].strip())
        set_run_font(run)
        return

    if stripped.endswith(":") and len(stripped) <= 60:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        run = paragraph.add_run(stripped)
        set_run_font(run, bold=True, color="1F4D78")
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(stripped)
    set_run_font(run)


def build_docx():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    skip_first_h1 = True
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if skip_first_h1 and stripped.startswith("# "):
            skip_first_h1 = False
            idx += 1
            continue
        skip_first_h1 = False

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue

        add_paragraph_from_text(doc, line)
        idx += 1

    add_footer(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
